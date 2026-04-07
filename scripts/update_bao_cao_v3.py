"""
Script cập nhật báo cáo lần 3:
  1. Xóa trang bìa thứ 2 (trùng lặp)
  2. Cập nhật Lời cam đoan và Lời cảm ơn → báo cáo nhóm, không cần chữ ký
  3. Xóa PHỤ LỤC (Phụ lục A - 72 features)
  4. Thay dấu "—" (em dash) thành "–" (en dash) toàn văn bản
"""
import sys, io, copy, shutil
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

SRC    = "dist/bao_cao_do_an_ids.docx"
BACKUP = f"dist/bao_cao_do_an_ids_backup_v3_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
OUT    = "dist/bao_cao_do_an_ids.docx"

EM  = "\u2014"   # — em dash (to replace)
EN  = "\u2013"   # – en dash (replacement)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def remove_body_child(body, idx):
    child = list(body)[idx]
    body.remove(child)
    return child


def body_children_snapshot(body):
    """Return list of (index, tag_short, first_text) for debug."""
    result = []
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        texts = [t.text or '' for t in child.iter(f'{{{W}}}t')]
        text = ''.join(texts).strip()[:50]
        result.append((i, tag, text))
    return result


def replace_em_dash_in_element(elem):
    """Replace all em dashes with en dashes in w:t text nodes."""
    for t in elem.iter(f'{{{W}}}t'):
        if t.text and EM in t.text:
            t.text = t.text.replace(EM, EN)
        if t.tail and EM in t.tail:
            t.tail = t.tail.replace(EM, EN)


def set_para_text(para, new_text):
    """Replace all runs with a single run containing new_text."""
    from docx.oxml import OxmlElement
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), 'Times New Roman')
    fonts.set(qn('w:hAnsi'), 'Times New Roman')
    fonts.set(qn('w:eastAsia'), 'Times New Roman')
    fonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(fonts)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = new_text
    r.append(t)
    p.append(r)


def find_para(doc, starts_with):
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts_with):
            return p
    return None


# ─── LOAD ──────────────────────────────────────────────────────────────────
print("Backing up...")
shutil.copy2(SRC, BACKUP)
print(f"  Backup: {BACKUP}")
doc = Document(SRC)
body = doc.element.body


# ===========================================================================
# 1. XÓA TRANG BÌA THỨ 2
#    Body children: [28]-[55] là trang bìa 2 + bảng GVHD/SVTH thứ 2
#    (xác định bằng cách tìm child chứa "TRƯỜNG ĐẠI HỌC" lần thứ 2)
# ===========================================================================
print("\n[1] Removing duplicate cover page...")

children = list(body)
# Find the 2nd occurrence of "TRƯỜNG ĐẠI HỌC" in body children
occurrences = []
for i, child in enumerate(children):
    texts = [t.text or '' for t in child.iter(f'{{{W}}}t')]
    text = ''.join(texts).strip()
    if 'TRƯỜNG ĐẠI HỌC CÔNG NGHỆ TP' in text:
        occurrences.append(i)

print(f"  'TRƯỜNG ĐẠI HỌC' found at body indices: {occurrences}")

if len(occurrences) >= 2:
    start_idx = occurrences[1]  # 2nd cover page starts here

    # Find end: the date paragraph of 2nd cover page
    # Go forward to find "TP. Hồ Chí Minh, tháng 4" after start_idx
    children = list(body)
    end_idx = None
    for i in range(start_idx, min(start_idx + 40, len(children))):
        child = children[i]
        texts = ''.join(t.text or '' for t in child.iter(f'{{{W}}}t')).strip()
        if 'TP. Hồ Chí Minh, tháng 4' in texts:
            end_idx = i
            break

    # Also remove the empty paragraph before the 2nd cover (page separator)
    separator_idx = start_idx - 1

    if end_idx is not None:
        print(f"  Removing body children [{separator_idx}..{end_idx}] (separator + 2nd cover page)")
        # Collect all to remove (must snapshot before removing, since indices shift)
        to_remove = list(body)[separator_idx:end_idx + 1]
        for child in to_remove:
            body.remove(child)
        print(f"  ✓ Removed {len(to_remove)} body children (2nd cover page)")
    else:
        print("  ✗ Could not find end of 2nd cover page")
else:
    print("  Only 1 cover page found, nothing to remove")


# ===========================================================================
# 2. CẬP NHẬT LỜI CAM ĐOAN → báo cáo nhóm, không cần chữ ký
# ===========================================================================
print("\n[2] Updating Lời cam đoan for group report...")

p72 = find_para(doc, 'Tôi xin cam đoan đây là công trình nghiên cứu của riêng tôi')
if p72:
    set_para_text(p72,
        'Nhóm xin cam đoan đây là công trình nghiên cứu của nhóm, được thực hiện dưới '
        'sự hướng dẫn tận tình của giảng viên hướng dẫn. Các số liệu, kết quả nêu trong '
        'báo cáo là trung thực, khách quan và chưa từng được công bố trong bất kỳ công '
        'trình nghiên cứu nào khác.')
    print("  ✓ Updated cam đoan paragraph 1")
else:
    print("  ✗ cam đoan paragraph 1 not found")

p74 = find_para(doc, 'Tôi xin chịu hoàn toàn trách nhiệm')
if p74:
    set_para_text(p74,
        'Nhóm xin chịu hoàn toàn trách nhiệm về tính trung thực của nội dung báo cáo này.')
    print("  ✓ Updated cam đoan paragraph 3")
else:
    print("  ✗ cam đoan paragraph 3 not found")

# Signature area — remove chữ ký, chỉ giữ ngày tháng + nhóm thực hiện
p75 = find_para(doc, 'TP. Hồ Chí Minh, Tháng 4 năm 2025\nSinh viên thực hiện')
if p75:
    set_para_text(p75, 'TP. Hồ Chí Minh, Tháng 4 năm 2025\n\nNhóm thực hiện')
    print("  ✓ Updated cam đoan signature area (removed individual signature)")
else:
    # Try alternative search
    for p in doc.paragraphs:
        if 'Sinh viên thực hiện' in p.text and 'Ký và ghi rõ' in p.text:
            set_para_text(p, 'TP. Hồ Chí Minh, Tháng 4 năm 2025\n\nNhóm thực hiện')
            print("  ✓ Updated cam đoan signature area (alt search)")
            break
    else:
        print("  ✗ cam đoan signature not found")


# ===========================================================================
# 3. CẬP NHẬT LỜI CẢM ƠN → dùng "nhóm" thay vì "tôi"
# ===========================================================================
print("\n[3] Updating Lời cảm ơn for group report...")

updates = [
    ('Trước tiên, tôi xin gửi lời cảm ơn chân thành',
     'Trước tiên, nhóm xin gửi lời cảm ơn chân thành và sâu sắc nhất đến giảng viên '
     'hướng dẫn, Thầy/Cô [TÊN GVHD], đã tận tình hướng dẫn, truyền đạt kiến thức quý '
     'báu và dành nhiều thời gian, tâm huyết để góp ý, định hướng giúp nhóm hoàn thành '
     'đồ án này.'),
    ('Tôi xin gửi lời cảm ơn sâu sắc đến quý thầy cô KHOA CÔNG NGHỆ',
     'Nhóm xin gửi lời cảm ơn sâu sắc đến quý thầy cô KHOA CÔNG NGHỆ THÔNG TIN, '
     'Trường Đại học Công nghệ TP. Hồ Chí Minh (HUTECH), đã trang bị cho nhóm nền '
     'tảng kiến thức vững chắc về an toàn thông tin, học máy và kỹ thuật hệ thống '
     'trong suốt quá trình học tập.'),
    ('Tôi cũng xin bày tỏ lòng biết ơn đối với Canadian Institute',
     'Nhóm cũng xin bày tỏ lòng biết ơn đối với Canadian Institute for Cybersecurity '
     '(CIC) thuộc University of New Brunswick đã công bố bộ dữ liệu CIC IoT-DIAD 2024 '
     '(nền tảng thực nghiệm quan trọng cho đồ án này), cũng như các cộng đồng mã nguồn '
     'mở đằng sau các công cụ scikit-learn, CatBoost, FastAPI và nhiều thư viện Python khác.'),
    ('Cuối cùng, tôi xin cảm ơn gia đình, bạn bè',
     'Cuối cùng, nhóm xin cảm ơn gia đình, bạn bè đã luôn động viên, hỗ trợ và tạo '
     'điều kiện thuận lợi để nhóm hoàn thành đồ án này.'),
    ('Mặc dù đã nỗ lực hết sức, báo cáo không thể tránh khỏi',
     'Mặc dù đã nỗ lực hết sức, báo cáo không thể tránh khỏi những thiếu sót. Nhóm '
     'rất mong nhận được những ý kiến đóng góp quý báu từ quý thầy cô và các bạn để '
     'báo cáo được hoàn thiện hơn.'),
]

for starts, new_text in updates:
    p = find_para(doc, starts)
    if p:
        set_para_text(p, new_text)
        print(f"  ✓ Updated: {starts[:50]}...")
    else:
        print(f"  ✗ Not found: {starts[:50]}...")


# ===========================================================================
# 4. XÓA PHỤ LỤC (Phụ lục A + bảng 72 features)
# ===========================================================================
print("\n[4] Removing PHỤ LỤC section...")

# Find "PHỤ LỤC" heading body child index
children = list(body)
phu_luc_idx = None
for i, child in enumerate(children):
    texts = ''.join(t.text or '' for t in child.iter(f'{{{W}}}t')).strip()
    if texts == 'PHỤ LỤC':
        phu_luc_idx = i
        break

if phu_luc_idx is not None:
    print(f"  Found PHỤ LỤC at body index {phu_luc_idx}")
    # Remove from phu_luc_idx to end of document
    children = list(body)
    to_remove = children[phu_luc_idx:]
    # But keep the final sectPr if present
    final_sectPr = None
    for child in reversed(to_remove):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            final_sectPr = child
            to_remove = to_remove[:-1]
            break
    for child in to_remove:
        body.remove(child)
    print(f"  ✓ Removed {len(to_remove)} elements (PHỤ LỤC section)")
else:
    print("  ✗ PHỤ LỤC heading not found")


# ===========================================================================
# 5. THAY DẤU "—" (em dash) → "–" (en dash) TOÀN VĂN BẢN
# ===========================================================================
print(f"\n[5] Replacing em dash '{EM}' with en dash '{EN}'...")

count = 0
# Replace in all paragraphs (including table cells)
replace_em_dash_in_element(body)

# Count replaced
all_text = ''.join(
    (t.text or '') for t in body.iter(f'{{{W}}}t')
)
remaining = all_text.count(EM)
# Count how many en dashes exist now
en_count = all_text.count(EN)
print(f"  ✓ Em dash replacement done")
print(f"    Remaining em dashes: {remaining}")
print(f"    En dashes in doc now: {en_count}")


# ===========================================================================
# 6. CẬP NHẬT BẢNG TRANG BÌA → "SVTH" thành "Nhóm SVTH" cho báo cáo nhóm
# ===========================================================================
print("\n[6] Updating cover page table for group report...")

table0 = doc.tables[0]
# Row 1 = SVTH row
svth_row = table0.rows[1]
label_cell = svth_row.cells[0]
for p in label_cell.paragraphs:
    if 'SVTH' in p.text:
        set_para_text(p, 'Nhóm SVTH:')
        print("  ✓ Updated SVTH → Nhóm SVTH")
        break

# Remove MSSV row (row 2) - for group, MSSV is N/A
# (Optional: comment out if you want to keep it)
# mssv_row = table0.rows[2]
# mssv_row._tr.getparent().remove(mssv_row._tr)
# print("  ✓ Removed MSSV row")


# ===========================================================================
# SAVE
# ===========================================================================
print(f"\nSaving to {OUT}...")
doc.save(OUT)
print("Done!")

print("\nSummary:")
print("  1. Duplicate cover page removed")
print("  2. Lời cam đoan updated for group (no signature)")
print("  3. Lời cảm ơn updated for group")
print("  4. PHỤ LỤC removed")
print(f" 5. Em dashes replaced with en dashes throughout")
print("  6. Cover page table: SVTH → Nhóm SVTH")
