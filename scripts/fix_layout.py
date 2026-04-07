import sys, io, shutil, copy
from datetime import datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches
from lxml import etree

SRC = "dist/bao_cao_do_an_ids.docx"
BACKUP = "dist/bao_cao_do_an_ids_backup_pre_layout_{}.docx".format(datetime.now().strftime('%Y%m%d_%H%M%S'))
shutil.copy2(SRC, BACKUP)
print("Backup: " + BACKUP)

doc = Document(SRC)
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def set_para_text(para, new_text):
    p_elem = para._p
    for r in p_elem.findall('{'+W+'}r'):
        p_elem.remove(r)
    for ins in p_elem.findall('{'+W+'}ins'):
        p_elem.remove(ins)
    r_elem = etree.SubElement(p_elem, '{'+W+'}r')
    t_elem = etree.SubElement(r_elem, '{'+W+'}t')
    t_elem.text = new_text
    if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def clone_para(doc, style_name, text):
    src = None
    for p in doc.paragraphs:
        if p.style.name == style_name and len(p.text.strip()) > 5:
            src = p
            break
    if src is None:
        raise RuntimeError("style not found: " + style_name)
    new_p = copy.deepcopy(src._p)
    for r in new_p.findall('{'+W+'}r'):
        new_p.remove(r)
    for ins in new_p.findall('{'+W+'}ins'):
        new_p.remove(ins)
    if text:
        r_elem = etree.SubElement(new_p, '{'+W+'}r')
        t_elem = etree.SubElement(r_elem, '{'+W+'}t')
        t_elem.text = text
        if text[0] == ' ' or text[-1] == ' ':
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# ── 1. Fix Bảng 5.4 caption: remove align=both ──────────────────────────────
fixed_54 = False
for p in doc.paragraphs:
    if p.style.name == 'TblCaption' and 'B\u1ea3ng 5.4' in p.text:
        pPr = p._p.find('{'+W+'}pPr')
        if pPr is not None:
            jc = pPr.find('{'+W+'}jc')
            if jc is not None:
                pPr.remove(jc)
                fixed_54 = True
                print("Fixed Bang 5.4 caption: removed align=both")
        break
if not fixed_54:
    print("WARNING: Bang 5.4 caption not found or no jc element")

# ── 2. Fix Hình 5.1 image width: 5.80in → 5.50in ────────────────────────────
EMU_TARGET = int(5.50 * 914400)  # 5.50 inches in EMU
fixed_img = False
for p in doc.paragraphs:
    extents = p._p.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
    )
    for ext in extents:
        cx = ext.get('cx')
        if cx and abs(int(cx) - int(5.80 * 914400)) < 50000:  # within ~0.05in tolerance
            old_cx = int(cx)
            old_cy = int(ext.get('cy', 0))
            # Scale height proportionally
            new_cx = EMU_TARGET
            new_cy = int(old_cy * new_cx / old_cx) if old_cx > 0 else old_cy
            ext.set('cx', str(new_cx))
            ext.set('cy', str(new_cy))
            # Also fix the <a:ext> inside <p:spPr> or drawingML
            parent = ext.getparent()
            if parent is not None:
                # Find sibling or child a:ext
                for sibling in parent.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}ext'):
                    sx = sibling.get('cx')
                    if sx and abs(int(sx) - old_cx) < 50000:
                        sibling.set('cx', str(new_cx))
                        sibling.set('cy', str(new_cy))
            fixed_img = True
            print("Fixed Hinh 5.1: {:.2f}in → {:.2f}in".format(old_cx/914400, new_cx/914400))
if not fixed_img:
    print("WARNING: 5.80in image not found")

# ── 3. Update console screen count and add missing screens ───────────────────
# Find "Console cung cap sau man hinh chinh"
p437 = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Console cung c\u1ea5p s\xe1u m\xe0n h\xecnh ch\xednh'):
        p437 = p
        break
if p437:
    set_para_text(p437,
        'Console cung c\u1ea5p ch\xedn m\xe0n h\xecnh ch\xednh cho workflow c\u1ee7a operator:'
    )
    print("Updated screen count text: sau -> chin")
else:
    print("WARNING: screen count para not found")

# Find the last bullet (Reports) and insert 3 new bullets after it
reports_bullet = None
for p in doc.paragraphs:
    if p.style.name == 'List Paragraph' and p.text.strip().startswith('Reports:'):
        reports_bullet = p
        break

if reports_bullet:
    new_bullets = [
        ('Live Logs: m\xe0n h\xecnh xem lu\u1ed3ng s\u1ef1 ki\u1ec7n live t\u1eeb sensor theo th\u1eddi gian th\u1ef1c, '
         'cho ph\xe9p operator theo d\xf5i ho\u1ea1t \u0111\u1ed9ng c\u1ee7a h\u1ec7 th\u1ed1ng m\xe0 kh\xf4ng c\u1ea7n \u0111\u1ecdc file log tr\u1ef1c ti\u1ebfp.'),
        ('Suppression Rules: m\xe0n h\xecnh qu\u1ea3n l\xfd c\xe1c quy t\u1eafc t\u1eaft ti\u1ebfng c\u1ea3nh b\xe1o cho c\xe1c lu\u1ed3ng \u0111\u01b0\u1ee3c x\xe1c nh\u1eadn l\xe0 l\xe0nh t\xednh, '
         'gi\xfap gi\u1ea3m false positive noise trong qu\xe1 tr\xecnh v\u1eadn h\xe0nh.'),
        ('System Health: m\xe0n h\xecnh ki\u1ec3m tra to\xe0n di\u1ec7n tr\u1ea1ng th\xe1i c\xe1c th\xe0nh ph\u1ea7n h\u1ec7 th\u1ed1ng giao di\u1ec7n tr\xfac quan '
         '(live sensor, inference engine, model bundle, notification worker).'),
    ]
    current_ref = reports_bullet._p
    for bullet_text in new_bullets:
        new_p = clone_para(doc, 'List Paragraph', bullet_text)
        current_ref.addnext(new_p)
        current_ref = new_p
    print("Added 3 new screen bullets (Live Logs, Suppression Rules, System Health)")
else:
    print("WARNING: Reports bullet not found")

doc.save(SRC)
print("\nSaved!")

# Verify
doc2 = Document(SRC)
print("\n=== Verify image widths ===")
for i, p in enumerate(doc2.paragraphs):
    extents = p._p.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
    )
    for ext in extents:
        cx = ext.get('cx')
        if cx:
            w = int(cx)/914400
            print("  [para {}] width={:.2f}in".format(i, w))

print("\n=== Verify Bang 5.4 caption ===")
for p in doc2.paragraphs:
    if 'B\u1ea3ng 5.4' in p.text and p.style.name == 'TblCaption':
        pPr = p._p.find('{'+W+'}pPr')
        jc = pPr.find('{'+W+'}jc') if pPr is not None else None
        print("  align jc element: {}".format(jc))

print("\n=== Verify console screen list ===")
in_section = False
for p in doc2.paragraphs:
    if 'Console cung c\u1ea5p' in p.text:
        in_section = True
    if in_section:
        print("  [{}]: {}".format(p.style.name, p.text[:100]))
    if in_section and p.style.name == 'Normal' and 'lane' in p.text:
        break
