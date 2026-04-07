"""
Script cập nhật báo cáo lần 2 - hoàn thiện báo cáo đầy đủ nhất.
Chạy từ root của project: python scripts/update_bao_cao_v2.py

Các thay đổi:
1. Thêm Bảng 5.4 (Stage 2 metrics table)
2. Cập nhật 5.6 (Alerts + Alert Detail)
3. Cập nhật 3.1 (Kiến trúc tổng thể)
4. Thêm 2.5 (Lý thuyết phân loại hai tầng)
5. Chèn hình kiến trúc (Hình 5.1) vào 5.7
6. Chèn biểu đồ F1 (Hình 5.2) vào 5.7.3
"""
import sys, io, copy, shutil, os
from datetime import datetime

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
import lxml.etree as etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SRC = "dist/bao_cao_do_an_ids.docx"
BACKUP = f"dist/bao_cao_do_an_ids_backup_v2_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
OUT = "dist/bao_cao_do_an_ids.docx"
IMG_ARCH = "/tmp/ids_two_stage_architecture.png"
IMG_F1   = "/tmp/ids_stage2_f1_chart.png"


# ─── XML helpers ────────────────────────────────────────────────────────────
def el(tag, attribs=None):
    e = OxmlElement(f'w:{tag}')
    if attribs:
        for k, v in attribs.items():
            e.set(qn(f'w:{k}'), v)
    return e


def make_run(text, bold=False):
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = el('rFonts', {'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman',
                          'eastAsia': 'Times New Roman', 'cs': 'Times New Roman'})
    rPr.append(fonts)
    if bold:
        rPr.append(el('b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def make_normal_para(text, bullet=False, center=False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if bullet:
        pPr.append(el('spacing', {'line': '360', 'lineRule': 'auto', 'after': '60'}))
        pPr.append(el('ind', {'left': '425', 'hanging': '283'}))
        pPr.append(el('jc', {'val': 'both'}))
    elif center:
        pPr.append(el('spacing', {'before': '0', 'after': '120', 'line': '360', 'lineRule': 'auto'}))
        pPr.append(el('jc', {'val': 'center'}))
    else:
        pPr.append(el('spacing', {'before': '0', 'after': '120', 'line': '360', 'lineRule': 'auto'}))
        pPr.append(el('ind', {'firstLine': '567'}))
        pPr.append(el('jc', {'val': 'both'}))
    p.append(pPr)
    full = ('• ' + text) if bullet else text
    p.append(make_run(full))
    return p


def make_figcaption(text):
    """Create a FigCaption-style paragraph matching document style."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = el('pStyle', {'val': 'FigCaption'})
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_tblcaption(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = el('pStyle', {'val': 'TblCaption'})
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_heading2(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pPr.append(el('pStyle', {'val': 'Heading2'}))
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_heading3(text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pPr.append(el('pStyle', {'val': 'Heading3'}))
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_after(ref_p, new_p):
    ref_p.addnext(new_p)


def find_para(doc, starts_with):
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts_with):
            return p
    return None


def set_para_text(para, new_text):
    p = para._p
    for r in p.findall(qn('w:r')):
        p.remove(r)
    p.append(make_run(new_text))


def make_picture_para(doc, img_path, width_inches=5.8):
    """Create a centered paragraph containing an inline picture."""
    # Add temp paragraph at end, grab its XML, then remove it from document
    tmp_para = doc.add_paragraph()
    run = tmp_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    # Make it centered
    tmp_para.alignment = 1  # CENTER
    p_elem = tmp_para._p
    # Detach from document body
    p_elem.getparent().remove(p_elem)
    return p_elem


def make_table_at_end(doc, headers, rows, col_widths_dxa):
    """Create a table at end of doc, detach it, return tbl element."""
    # Calculate total width
    total = sum(col_widths_dxa)
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'

    # Set column widths
    for i, row in enumerate(tbl.rows):
        for j, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = el('tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(col_widths_dxa[j]))
            tcW.set(qn('w:type'), 'dxa')

    # Header row - bold
    header_row = tbl.rows[0]
    for j, hdr in enumerate(headers):
        cell = header_row.cells[j]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = 1
        run = para.add_run(hdr)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

    # Data rows
    for i, data_row in enumerate(rows):
        row = tbl.rows[i + 1]
        for j, val in enumerate(data_row):
            cell = row.cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = 1 if j > 0 else 0
            run = para.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    tbl_elem = tbl._tbl
    tbl_elem.getparent().remove(tbl_elem)
    return tbl_elem


# ─── MAIN ───────────────────────────────────────────────────────────────────
print("Backing up...")
shutil.copy2(SRC, BACKUP)
print(f"  Backup: {BACKUP}")

print("Loading document...")
doc = Document(SRC)

# ===========================================================================
# 1. BẢNG 5.4 — Stage 2 per-family metrics table
# ===========================================================================
print("\n[1] Adding Bảng 5.4 (Stage 2 metrics)...")

# The caption paragraph "Bảng 5.4 tổng hợp..." already exists.
# Replace it with a proper TblCaption style, then insert the actual table after it.
caption_para = find_para(doc, 'Bảng 5.4 tổng hợp kết quả của Stage 2')
if caption_para:
    # Re-style the caption paragraph to TblCaption
    pPr = caption_para._p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        caption_para._p.insert(0, pPr)
    # Set style to TblCaption
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = el('pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), 'TblCaption')
    # Update caption text
    set_para_text(caption_para,
        'Bảng 5.4. Kết quả Stage 2 family classifier trên tập test (full-data, 18.457.115 attack rows)')

    # Create the table
    headers = ['Họ tấn công', 'Phân loại', 'F1-score', 'Ghi chú']
    rows = [
        ['DoS',        'In-distribution', '0,9935', 'Họ đa số, hiệu năng xuất sắc'],
        ['DDoS',       'In-distribution', '0,9338', 'Họ đa số, hiệu năng tốt'],
        ['Mirai',      'In-distribution', '0,4488', 'Họ trung bình, cần cải thiện'],
        ['Spoofing',   'In-distribution', '0,2955', 'Họ thiểu số, F1 thấp'],
        ['Web-Based',  'In-distribution', '0,0162', 'Họ rất thiểu số, gần như không học được'],
        ['BruteForce', 'OOD (probe)',     '—',      'Giữ lại làm OOD probe, không trong closed-set'],
        ['Recon',      'OOD (probe)',     '—',      'Giữ lại làm OOD probe, không trong closed-set'],
        ['Weighted F1',     'Tổng hợp', '0,9804', 'Bị kéo lên bởi DoS/DDoS đa số'],
        ['Macro F1',        'Tổng hợp', '0,5376', 'Phản ánh mất cân bằng giữa các họ'],
        ['Accuracy',        'Tổng hợp', '0,9775', 'Trên tập test'],
    ]
    # col widths in DXA (total ~9360 for letter with 1in margins)
    col_w = [1800, 1800, 1200, 4560]
    tbl_elem = make_table_at_end(doc, headers, rows, col_w)
    insert_after(caption_para._p, tbl_elem)
    print("  ✓ Bảng 5.4 added")
else:
    print("  ✗ Could not find Bảng 5.4 caption")

# ===========================================================================
# 2. INSERT ARCHITECTURE IMAGE (Hình 5.1) — after Section 5.7 intro paragraph
# ===========================================================================
print("\n[2] Inserting Hình 5.1 (architecture diagram)...")

intro_57 = find_para(doc, 'Hệ thống phân loại hai tầng là thành phần mở rộng được bổ sung')
if intro_57:
    img_p = make_picture_para(doc, IMG_ARCH, width_inches=5.8)
    cap_p = make_figcaption(
        'Hình 5.1. Kiến trúc hệ thống phân loại hai tầng\n'
        'Nguồn: Thiết kế của tác giả')
    # Insert caption first (after intro), then image (after intro) → image ends up before caption
    insert_after(intro_57._p, cap_p)
    insert_after(intro_57._p, img_p)
    print("  ✓ Hình 5.1 inserted")
else:
    print("  ✗ Could not find 5.7 intro paragraph")

# ===========================================================================
# 3. INSERT F1 BAR CHART (Hình 5.2) — after the results paragraph in 5.7.3
# ===========================================================================
print("\n[3] Inserting Hình 5.2 (F1 bar chart)...")

results_para = find_para(doc,
    'Stage 2 đạt Weighted F1 = 0,9804 và Accuracy = 0,9775')
if results_para:
    img_p2 = make_picture_para(doc, IMG_F1, width_inches=5.5)
    cap_p2 = make_figcaption(
        'Hình 5.2. F1-score của Stage 2 theo từng họ tấn công\n'
        'Nguồn: Kết quả thực nghiệm của tác giả')
    insert_after(results_para._p, cap_p2)
    insert_after(results_para._p, img_p2)
    print("  ✓ Hình 5.2 inserted")
else:
    print("  ✗ Could not find Stage 2 results paragraph")

# ===========================================================================
# 4. UPDATE SECTION 5.6 — Alerts and Alert Detail screen descriptions
# ===========================================================================
print("\n[4] Updating Section 5.6 screen descriptions...")

alerts_bullet = find_para(doc,
    '• Alerts: lane làm việc chính cho analyst, hiển thị danh sách các cảnh báo')
if alerts_bullet:
    set_para_text(alerts_bullet,
        '• Alerts: lane làm việc chính cho analyst, hiển thị danh sách các cảnh báo '
        'attack kèm filter, sort và drill-down. Sau khi tích hợp Stage 2, mỗi hàng '
        'trong danh sách hiển thị thêm compact family badge cho thấy attack_family '
        '(DDoS, DoS, Mirai, Spoofing, Web-Based) và family_status (known/unknown); '
        'operator có thể nhận biết họ tấn công ngay khi scan qua danh sách mà không '
        'cần mở từng alert.')
    print("  ✓ Updated Alerts bullet")
else:
    print("  ✗ Could not find Alerts bullet")

detail_bullet = find_para(doc,
    '• Alert Detail: trang chi tiết cho từng cảnh báo, hiển thị đầy đủ 72 feature')
if detail_bullet:
    set_para_text(detail_bullet,
        '• Alert Detail: trang chi tiết cho từng cảnh báo, hiển thị đầy đủ 72 feature, '
        'attack_score, và các thông tin truy vết. Sau tích hợp Stage 2, trang này bổ '
        'sung thêm phần family enrichment: attack_family, family_status, '
        'attack_family_confidence, attack_family_margin, và ngữ cảnh abstention '
        '(nếu family_status = "unknown"). Operator có thể hiểu vì sao một cảnh báo '
        'được phân loại "known" hay "unknown" mà không cần đọc log kỹ thuật.')
    print("  ✓ Updated Alert Detail bullet")
else:
    print("  ✗ Could not find Alert Detail bullet")

# ===========================================================================
# 5. UPDATE SECTION 3.1 — Add two-stage mention in Tầng 4
# ===========================================================================
print("\n[5] Updating Section 3.1 — Tầng 4 description...")

tang4_para = find_para(doc,
    'Tầng 4 (Runtime inference): nhận các structured flow records từ upstream')
if tang4_para:
    set_para_text(tang4_para,
        'Tầng 4 (Runtime inference): nhận các structured flow records từ upstream, '
        'kiểm tra contract 72 đặc trưng, chạy model dự đoán và trả về kết quả '
        'alert/benign cho downstream. Sau khi tích hợp hệ thống hai tầng, Tầng 4 '
        'thực hiện tuần tự: Stage 1 binary inference (Benign/Attack), '
        'và nếu is_alert = True thì chạy tiếp Stage 2 family inference '
        '(DDoS/DoS/Mirai/Spoofing/Web-Based hoặc unknown). '
        'Toàn bộ chuỗi inference này được đóng gói trong composite bundle '
        'và resolve thông qua activation record host-local duy nhất.')
    print("  ✓ Updated Tầng 4 description")
else:
    print("  ✗ Could not find Tầng 4 paragraph")

# Also add Tầng 5 mention about console family surfaces
tang5_para = find_para(doc,
    'Tầng 5 (Operator console): cung cấp giao diện web cho operator giám sát')
if tang5_para:
    set_para_text(tang5_para,
        'Tầng 5 (Operator console): cung cấp giao diện web cho operator giám sát '
        'cảnh báo, tình trạng hệ thống và lịch sử. Stack kỹ thuật: FastAPI + Jinja2 '
        '+ SQLite, chạy same-host, đặt sau NGINX reverse proxy. '
        'Sau tích hợp Stage 2, console hiển thị thêm thông tin family enrichment '
        '(attack_family, family_status) trên màn hình Alerts và Alert Detail, '
        'cho phép operator phân loại và ưu tiên xử lý theo họ tấn công.')
    print("  ✓ Updated Tầng 5 description")
else:
    print("  ✗ Could not find Tầng 5 paragraph")

# ===========================================================================
# 6. ADD SECTION 2.5 — Theory for two-stage classification
# ===========================================================================
print("\n[6] Adding Section 2.5 — Lý thuyết phân loại hai tầng...")

# Insert before CHƯƠNG 3 (find last paragraph of Chapter 2)
chuong3_heading = find_para(doc, 'CHƯƠNG 3. PHƯƠNG PHÁP VÀ THIẾT KẾ HỆ THỐNG')
if chuong3_heading:
    # Build paragraphs in reverse order (each inserted immediately before Chương 3)
    paras_25 = []

    paras_25.append(make_heading2('2.5 Kiến trúc phân loại hai tầng và vấn đề lớp mở'))

    paras_25.append(make_normal_para(
        'Phân loại đa lớp (multi-class classification) mở rộng bài toán nhị phân '
        'bằng cách dự đoán một nhãn từ không gian nhãn có nhiều hơn hai lớp. '
        'Trong bối cảnh IDS, phân loại đa lớp cho phép xác định không chỉ '
        '"có tấn công hay không" mà còn "đây là loại tấn công gì", cung cấp '
        'thông tin có giá trị cao hơn cho SOC analyst khi ứng phó sự cố. '
        'Tuy nhiên, các IDS thực tế phải đối mặt với vấn đề lớp mở (open-world problem): '
        'không gian tấn công thực tế là vô hạn và luôn xuất hiện các loại tấn công '
        'mới chưa từng thấy trong dữ liệu huấn luyện [10].'))

    paras_25.append(make_normal_para(
        'Phân loại hai tầng (two-stage classification hay cascaded classifier) là '
        'một kiến trúc giải quyết đồng thời hai vấn đề trên. Tầng thứ nhất '
        '(coarse classifier) phân loại vào một trong hai nhóm lớn (ví dụ: Benign/Attack), '
        'đóng vai trò cổng lọc sơ bộ. Tầng thứ hai (fine-grained classifier) '
        'chỉ chạy trên các mẫu vượt qua cổng tầng 1, thực hiện phân loại chi tiết '
        'hơn trong nhóm đó (ví dụ: phân loại họ tấn công). '
        'Ưu điểm của kiến trúc này là: (1) cho phép tái sử dụng và giữ nguyên mô hình '
        'tầng 1 đã được kiểm chứng; (2) tầng 2 chỉ học trên không gian con có phân phối '
        'đồng nhất hơn, giúp tăng độ chính xác; (3) dễ dàng thêm tầng 2 vào hệ thống '
        'hiện có mà không phá vỡ contract đã có.'))

    paras_25.append(make_normal_para(
        'Cơ chế từ chối/bỏ phiếu trắng (abstention hay rejection option) là kỹ thuật '
        'cho phép mô hình phân loại "không chắc" thay vì buộc phải ép một mẫu vào '
        'lớp gần nhất trong tập đóng. Trong bài toán IDS với lớp mở, abstention '
        'đặc biệt quan trọng: khi gặp một mẫu tấn công thuộc họ chưa được huấn luyện '
        '(out-of-distribution — OOD), mô hình cần thừa nhận sự không chắc chắn thay vì '
        'nhãn nhầm thành một họ đã biết. '
        'Abstention thường được triển khai bằng cách đặt ngưỡng trên xác suất top-1 '
        '(confidence threshold) hoặc kết hợp thêm điều kiện về khoảng cách giữa '
        'top-1 và top-2 (margin threshold). '
        'Nếu mẫu không vượt qua ngưỡng, mô hình trả về nhãn "unknown/abstain" '
        'thay vì nhãn lớp cụ thể [11].'))

    paras_25.append(make_normal_para(
        'Trong đồ án này, kiến trúc hai tầng được triển khai cụ thể như sau: '
        'Stage 1 là binary CatBoost classifier (Benign/Attack) đã được kiểm chứng '
        'qua bốn vòng thực nghiệm ở Chương 4; Stage 2 là CatBoost multiclass classifier '
        'với closed-set gồm năm họ (DDoS, DoS, Mirai, Spoofing, Web-Based) và cơ chế '
        'abstain hai tín hiệu (top1_confidence và runner_up_margin) được calibrate '
        'trên tập validation. Abstention được mã hóa trong trường family_status = "unknown", '
        'phân biệt rõ với "benign" (Stage 1 không phát hiện tấn công) và '
        '"known" (Stage 2 đủ tự tin gán họ cụ thể).'))

    # Insert in reverse order so they appear in correct sequence
    for p_elem in reversed(paras_25):
        insert_after(chuong3_heading._p.getprevious(), p_elem)

    print(f"  ✓ Section 2.5 inserted ({len(paras_25)} paragraphs)")
else:
    print("  ✗ Could not find CHƯƠNG 3 heading")

# ===========================================================================
# 7. SAVE
# ===========================================================================
print(f"\nSaving to {OUT}...")
doc.save(OUT)
print("Done!")

print("\nSummary of changes:")
print("  1. Bảng 5.4 — Stage 2 per-family metrics table added")
print("  2. Hình 5.1 — Two-stage architecture diagram inserted")
print("  3. Hình 5.2 — Stage 2 F1 bar chart inserted")
print("  4. Section 5.6 — Alerts + Alert Detail updated with family enrichment")
print("  5. Section 3.1 — Tầng 4 + Tầng 5 descriptions updated")
print("  6. Section 2.5 — Two-stage theory added (4 paragraphs)")
