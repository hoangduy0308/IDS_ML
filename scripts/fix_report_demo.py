import sys, io, shutil, copy
from datetime import datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from lxml import etree

SRC = "dist/bao_cao_do_an_ids.docx"
BACKUP = "dist/bao_cao_do_an_ids_backup_pre_demo_{}.docx".format(datetime.now().strftime('%Y%m%d_%H%M%S'))
shutil.copy2(SRC, BACKUP)
print("Backup: " + BACKUP)

doc = Document(SRC)
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def set_para_text(para, new_text):
    p_elem = para._p
    for r in p_elem.findall('{'+W+'}r'):
        p_elem.remove(r)
    for ins in p_elem.findall('{'+W+'}ins'):
        p_elem.remove(ins)
    r_elem = etree.SubElement(p_elem, '{'+W+'}r')
    t_elem = etree.SubElement(r_elem, '{'+W+'}t')
    t_elem.text = new_text
    if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
        t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

def clone_para(doc, style_name, text):
    src = None
    for p in doc.paragraphs:
        if p.style.name == style_name and len(p.text.strip()) > 5:
            src = p
            break
    if src is None:
        raise RuntimeError("style not found: " + style_name)
    new_p = copy.deepcopy(src._p)
    for r in new_p.findall('{'+W+'}r'):
        new_p.remove(r)
    for ins in new_p.findall('{'+W+'}ins'):
        new_p.remove(ins)
    if text:
        r_elem = etree.SubElement(new_p, '{'+W+'}r')
        t_elem = etree.SubElement(r_elem, '{'+W+'}t')
        t_elem.text = text
        if text[0] == ' ' or text[-1] == ' ':
            t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return new_p

# ── 1. Fix Bang 4.2 OS ──────────────────────────────────────────────────────
for t in doc.tables:
    rows = [[c.text.strip() for c in r.cells] for r in t.rows]
    if any('Windows 11' in str(r) for r in rows):
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells[0] == 'OS':
                row.cells[1].paragraphs[0].clear()
                row.cells[1].paragraphs[0].add_run(
                    'Ubuntu 22.04 LTS / Kali Linux (m\u00f4i tr\u01b0\u1eddng tri\u1ec3n khai demo)')
                print("Fixed Bang 4.2 OS")
            elif cells[0] == 'Python':
                row.cells[1].paragraphs[0].clear()
                row.cells[1].paragraphs[0].add_run('3.11.x')
                print("Fixed Bang 4.2 Python")
        break

# ── 2. Update para 326 ───────────────────────────────────────────────────────
p326 = doc.paragraphs[326]
p326_text = (
    'Vi\u1ec7c ki\u1ec3m th\u1eed inference engine, realtime pipeline v\xe0 c\xe1c th\xe0nh ph\u1ea7n runtime '
    '\u0111\u01b0\u1ee3c th\u1ef1c hi\u1ec7n tr\xean m\xe1y tri\u1ec3n khai Linux '
    '(Ubuntu 22.04 LTS / Kali Linux) v\u1edbi c\u1ea5u h\xecnh sau. '
    'M\xf4i tr\u01b0\u1eddng n\xe0y c\u0169ng l\xe0 m\xf4i tr\u01b0\u1eddng ch\u1ea1y demo th\u1ef1c t\u1ebf c\u1ee7a h\u1ec7 th\u1ed1ng.'
)
set_para_text(p326, p326_text)
print("Updated para 326")

# ── 3. Add Telegram para after operator console section ─────────────────────
tele_ref = None
for p in doc.paragraphs:
    if 'V\xf2ng \u0111\u1eddi schema' in p.text and 'console' in p.text.lower():
        tele_ref = p
        break
if tele_ref is None:
    for i, p in enumerate(doc.paragraphs):
        if 'H\u1ec7 th\u1ed1ng ph\xe2n lo\u1ea1i hai t\u1ea7ng' in p.text and p.style.name.startswith('Heading'):
            tele_ref = doc.paragraphs[i-1]
            break
print("Telegram ref: " + (tele_ref.text[:60] if tele_ref else "NOT FOUND"))

if tele_ref:
    tele_text = (
        'Ngo\xe0i giao di\u1ec7n web, h\u1ec7 th\u1ed1ng c\xf2n t\xedch h\u1ee3p k\xeanh th\xf4ng b\xe1o ch\u1ee7 \u0111\u1ed9ng qua Telegram bot. '
        'Khi c\xf3 c\u1ea3nh b\xe1o m\u1edbi, notification worker g\u1eedi tin nh\u1eafn t\u1ee9c th\xec \u0111\u1ebfn operator qua API Telegram v\u1edbi hai d\u1ea1ng: '
        '[IDS ALERT] cho t\u1eebng c\u1ea3nh b\xe1o \u0111\u01a1n l\u1ebb (k\xe8m event id, flow, label, family, score, timestamp) v\xe0 '
        '[IDS ALERT BURST] cho nh\xf3m c\u1ea3nh b\xe1o c\xf9ng lu\u1ed3ng (k\xe8m incident id, s\u1ed1 alerts, flow t\u1ed5ng h\u1ee3p, score_max, dst_ports). '
        'T\xednh n\u0103ng n\xe0y cho ph\xe9p operator gi\xe1m s\xe1t IDS t\u1eeb xa m\xe0 kh\xf4ng c\u1ea7n lu\xf4n m\u1edf tr\xecnh duy\u1ec7t.'
    )
    new_tele_p = clone_para(doc, 'Normal', tele_text)
    tele_ref._p.addnext(new_tele_p)
    print("Added Telegram para")

# ── 4. Add Section 6.5 Demo ─────────────────────────────────────────────────
tltk_para = None
for p in doc.paragraphs:
    if p.text.strip() == 'T\xc0I LI\u1ec6U THAM KH\u1ea2O':
        tltk_para = p
        break
print("TLTK: " + (tltk_para.text if tltk_para else "NOT FOUND"))

if tltk_para:
    ref = tltk_para._p

    blocks = []
    blocks.append(clone_para(doc, 'Heading 2', '6.5 Demo h\u1ec7 th\u1ed1ng th\u1ef1c t\u1ebf'))
    blocks.append(clone_para(doc, 'Normal',
        'Ph\u1ea7n n\xe0y tr\xecnh b\xe0y k\u1ebft qu\u1ea3 demo h\u1ec7 th\u1ed1ng IDS VIGIL \u0111ang ch\u1ea1y tr\xean m\xf4i tr\u01b0\u1eddng Linux th\u1ef1c t\u1ebf, '
        'bao g\u1ed3m operator console web v\xe0 k\xeanh th\xf4ng b\xe1o Telegram. Demo \u0111\u01b0\u1ee3c th\u1ef1c hi\u1ec7n v\u1edbi traffic m\u1ea1ng th\u1ef1c t\u1ebf '
        'trong m\xf4i tr\u01b0\u1eddng ki\u1ec3m th\u1eed c\xf3 ki\u1ec3m so\xe1t, trong \u0111\xf3 live sensor thu th\u1eadp v\xe0 ph\xe2n lo\u1ea1i traffic theo th\u1eddi gian th\u1ef1c.'
    ))
    blocks.append(clone_para(doc, 'Heading 3', '6.5.1 M\xe0n h\xecnh Overview v\xe0 tr\u1ea1ng th\xe1i h\u1ec7 th\u1ed1ng'))
    blocks.append(clone_para(doc, 'Normal',
        'M\xe0n h\xecnh Overview (H\xecnh 6.1) hi\u1ec3n th\u1ecb b\u1ed1n ch\u1ec9 s\u1ed1 t\u1ed5ng quan: t\u1ed5ng s\u1ed1 c\u1ea3nh b\xe1o, s\u1ed1 anomaly \u0111ang active, '
        'system readiness v\xe0 system status. Trong phi\xean demo, h\u1ec7 th\u1ed1ng \u0111\u1ea1t System Readiness 100% (4/4 components OK) '
        'v\u1edbi tr\u1ea1ng th\xe1i "Ready \u2013 operational". Sidebar Alert Triage ph\xe2n lo\u1ea1i c\u1ea3nh b\xe1o theo tr\u1ea1ng th\xe1i '
        'x\u1eed l\xfd: New, Acknowledged, Investigating, Resolved, False Positive. '
        'Recent Alerts hi\u1ec3n th\u1ecb c\u1ea3nh b\xe1o m\u1edbi nh\u1ea5t v\u1edbi grouping theo lu\u1ed3ng src IP \u2192 dst IP.'
    ))
    blocks.append(clone_para(doc, 'Normal', ''))
    blocks.append(clone_para(doc, 'FigCaption',
        'H\xecnh 6.1. M\xe0n h\xecnh Overview c\u1ee7a VIGIL IDS \u2013 System Readiness 100%, System Status: Ready operational'
    ))
    blocks.append(clone_para(doc, 'Heading 3', '6.5.2 M\xe0n h\xecnh Alerts v\xe0 ph\xe2n lo\u1ea1i h\u1ecd t\u1ea5n c\xf4ng'))
    blocks.append(clone_para(doc, 'Normal',
        'M\xe0n h\xecnh Alerts (H\xecnh 6.2) l\xe0 giao di\u1ec7n l\xe0m vi\u1ec7c ch\xednh cho analyst. Trong phi\xean demo v\u1edbi 12 c\u1ea3nh b\xe1o, '
        'h\u1ec7 th\u1ed1ng ph\xe2n lo\u1ea1i ch\xednh x\xe1c c\xe1c h\u1ecd t\u1ea5n c\xf4ng qua Stage 2: DoS \u0111\u01b0\u1ee3c nh\u1eadn di\u1ec7n t\u1eeb lu\u1ed3ng UDP '
        '192.168.117.128 \u2192 192.168.117.2 v\xe0 lu\u1ed3ng TCP 149.154.166.110, Mirai t\u1eeb lu\u1ed3ng TCP \u0111\u1ebfn 149.154.166.110. '
        'M\u1ed9t s\u1ed1 c\u1ea3nh b\xe1o hi\u1ec3n th\u1ecb "unknown family" \u2013 c\u01a1 ch\u1ebf abstain ho\u1ea1t \u0111\u1ed9ng \u0111\xfang: Stage 1 x\xe1c nh\u1eadn t\u1ea5n c\xf4ng '
        'nh\u01b0ng Stage 2 kh\xf4ng \u0111\u1ee7 confidence \u0111\u1ec3 g\xe1n h\u1ecd c\u1ee5 th\u1ec3. C\u1ed9t FAMILY SIGNAL ph\xe2n bi\u1ec7t '
        '"known family" v\xe0 "unknown family", gi\xfap operator triage nhanh m\xe0 kh\xf4ng c\u1ea7n m\u1edf t\u1eebng alert.'
    ))
    blocks.append(clone_para(doc, 'Normal', ''))
    blocks.append(clone_para(doc, 'FigCaption',
        'H\xecnh 6.2. M\xe0n h\xecnh Alerts c\u1ee7a VIGIL IDS \u2013 hi\u1ec3n th\u1ecb Family Signal (known/unknown) v\xe0 12 c\u1ea3nh b\xe1o ph\xe2n lo\u1ea1i'
    ))
    blocks.append(clone_para(doc, 'Heading 3', '6.5.3 Th\xf4ng b\xe1o Telegram realtime'))
    blocks.append(clone_para(doc, 'Normal',
        'K\xeanh th\xf4ng b\xe1o Telegram (H\xecnh 6.3) cho th\u1ea5y notification worker g\u1eedi c\u1ea3nh b\xe1o t\u1ee9c th\xec khi c\xf3 attack m\u1edbi. '
        'Hai d\u1ea1ng tin nh\u1eafn xu\u1ea5t hi\u1ec7n trong phi\xean demo: [IDS ALERT] cho c\u1ea3nh b\xe1o \u0111\u01a1n l\u1ebb v\u1edbi \u0111\u1ea7y \u0111\u1ee7 th\xf4ng tin '
        'flow (src:port \u2192 dst:port), protocol, label, family (DoS/unknown), attack_score; v\xe0 [IDS ALERT BURST] '
        'cho nh\xf3m c\u1ea3nh b\xe1o c\xf9ng lu\u1ed3ng v\u1edbi incident_id, t\u1ed5ng s\u1ed1 alerts, score_max v\xe0 dst_ports t\u1ed5ng h\u1ee3p. '
        'Th\xf4ng b\xe1o \u0111\u01b0\u1ee3c g\u1eedi trong v\xf2ng v\xe0i gi\xe2y sau khi live sensor ph\xe1t hi\u1ec7n attack, '
        '\u0111\u1ea3m b\u1ea3o operator c\xf3 th\u1ec3 nh\u1eadn alert ngay c\u1ea3 khi kh\xf4ng \u0111ang m\u1edf operator console.'
    ))
    blocks.append(clone_para(doc, 'Normal', ''))
    blocks.append(clone_para(doc, 'FigCaption',
        'H\xecnh 6.3. Th\xf4ng b\xe1o Telegram realtime t\u1eeb VIGIL IDS bot \u2013 [IDS ALERT] v\xe0 [IDS ALERT BURST]'
    ))

    current_ref = ref
    for block in reversed(blocks):
        current_ref.addprevious(block)
        current_ref = block

    print("Inserted {} blocks for Section 6.5".format(len(blocks)))

doc.save(SRC)
print("Saved!")

doc2 = Document(SRC)
print("\n=== Verify ===")
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if '6.5' in t or 'H\xecnh 6.' in t or 'VIGIL' in t[:30]:
        print("[{}] [{}]: '{}'".format(i, p.style.name, t[:100]))
