import sys, io, shutil, copy
from datetime import datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

SRC = "dist/bao_cao_do_an_ids.docx"
BACKUP = "dist/bao_cao_do_an_ids_backup_pre_images_{}.docx".format(datetime.now().strftime('%Y%m%d_%H%M%S'))
shutil.copy2(SRC, BACKUP)
print("Backup: " + BACKUP)

IMAGES = [
    (r"C:\Users\hdi\Downloads\z7695155318779_efe6685147d14f6be017ca96ca0a9017.jpg", "6.1"),
    (r"C:\Users\hdi\Downloads\z7695156469984_f71f937516a2bc6df75355ade4cd115d.jpg", "6.2"),
    (r"C:\Users\hdi\Downloads\z7695155591433_e5b175f32baaa0a66b4f1c1f04cf6c37.jpg", "6.3"),
]

doc = Document(SRC)
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def make_picture_para(doc, img_path, width_inches=5.5):
    """Create a centered paragraph with an image, detach and return _p element."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    p_elem = p._p
    p_elem.getparent().remove(p_elem)
    return p_elem

# Find each blank placeholder before the FigCaption for Hình 6.x
# Structure: [...text...] [blank] [FigCaption "Hình 6.x..."]
# We replace the blank with the actual image

for img_path, fig_num in IMAGES:
    cap_text = "H\xecnh " + fig_num + "."
    cap_para = None
    for p in doc.paragraphs:
        if p.style.name == 'FigCaption' and p.text.strip().startswith(cap_text):
            cap_para = p
            break

    if cap_para is None:
        print("WARNING: FigCaption for {} not found".format(fig_num))
        continue

    print("Found caption for {}: '{}'".format(fig_num, cap_para.text[:60]))

    # Find the blank para immediately before this caption in the XML body
    # Walk siblings backwards from cap_para._p
    prev_elem = cap_para._p.getprevious()

    # Check if it's a blank/empty paragraph
    if prev_elem is not None:
        prev_texts = ''.join(t.text or '' for t in prev_elem.iter('{'+W+'}t')).strip()
        print("  Prev elem text: '{}'".format(prev_texts[:40]))

        # Create image paragraph
        img_p = make_picture_para(doc, img_path, width_inches=5.5)

        if prev_texts == '':
            # Replace the blank para with the image para
            prev_elem.getparent().replace(prev_elem, img_p)
            print("  Replaced blank para with image for Hinh {}".format(fig_num))
        else:
            # Insert image before caption
            cap_para._p.addprevious(img_p)
            print("  Inserted image before caption for Hinh {}".format(fig_num))

doc.save(SRC)
print("\nSaved!")

# Verify figure count
doc2 = Document(SRC)
fig_caps = [p for p in doc2.paragraphs if p.style.name == 'FigCaption']
print("\nAll figure captions:")
for p in fig_caps:
    print("  " + p.text[:80])
