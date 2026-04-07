"""
Script cập nhật báo cáo đồ án với nội dung hệ thống phân loại hai tầng (two-stage).
Chạy từ root của project: python scripts/update_bao_cao.py
"""
import sys
import io
import copy
import shutil
import os
from datetime import datetime

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

SRC = "dist/bao_cao_do_an_ids.docx"
BACKUP = f"dist/bao_cao_do_an_ids_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
OUT = "dist/bao_cao_do_an_ids.docx"


def el(tag, attribs=None, text=None):
    """Create a w: element with optional attribs and text."""
    e = OxmlElement(f'w:{tag}')
    if attribs:
        for k, v in attribs.items():
            e.set(qn(f'w:{k}'), v)
    if text is not None:
        e.text = text
    return e


def make_run(text, bold=False, font_size=None, font_name="Times New Roman"):
    """Create a w:r run element with given text and optional formatting."""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    fonts = el('rFonts', {'ascii': font_name, 'hAnsi': font_name, 'eastAsia': font_name, 'cs': font_name})
    rPr.append(fonts)
    if bold:
        rPr.append(el('b'))
    if font_size:
        rPr.append(el('sz', {'val': str(font_size * 2)}))
        rPr.append(el('szCs', {'val': str(font_size * 2)}))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    return r


def make_normal_para(text, bullet=False):
    """Create a Normal paragraph (with or without bullet) matching document style."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    if bullet:
        # Match bullet paragraph pPr from document
        spacing = el('spacing', {'line': '360', 'lineRule': 'auto', 'after': '60'})
        ind = el('ind', {'left': '425', 'hanging': '283'})
        jc = el('jc', {'val': 'both'})
        pPr.append(spacing)
        pPr.append(ind)
        pPr.append(jc)
    else:
        # Match normal paragraph pPr from document
        spacing = el('spacing', {'before': '0', 'after': '120', 'line': '360', 'lineRule': 'auto'})
        ind = el('ind', {'firstLine': '567'})
        jc = el('jc', {'val': 'both'})
        pPr.append(spacing)
        pPr.append(ind)
        pPr.append(jc)

    p.append(pPr)
    full_text = ('• ' + text) if bullet else text
    p.append(make_run(full_text))
    return p


def make_heading2(text):
    """Create a Heading 2 paragraph."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = el('pStyle', {'val': 'Heading2'})
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def make_heading3(text):
    """Create a Heading 3 paragraph."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = el('pStyle', {'val': 'Heading3'})
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def insert_after(ref_p_elem, new_p_elem):
    """Insert new_p_elem immediately after ref_p_elem in the document body."""
    ref_p_elem.addnext(new_p_elem)


def find_para_by_text(doc, starts_with):
    """Find a paragraph that starts with the given text string."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(starts_with):
            return p
    return None


def get_para_text(doc, idx):
    return doc.paragraphs[idx].text.strip()


def set_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run containing new_text."""
    p = para._p
    # Remove all existing runs
    for r in p.findall(qn('w:r')):
        p.remove(r)
    # Add new run
    p.append(make_run(new_text))


def add_table_row(table, cell_texts):
    """Add a new row to table with given cell texts, copying style from last row."""
    last_row = table.rows[-1]
    new_tr = copy.deepcopy(last_row._tr)
    # Clear text in cells
    cells = new_tr.findall(f'.//{{{W}}}tc')
    for i, tc in enumerate(cells):
        # Remove all paragraphs from cell
        for p in tc.findall(f'{{{W}}}p'):
            tc.remove(p)
        # Add new paragraph with text
        new_p = OxmlElement('w:p')
        # Copy paragraph properties from original cell
        orig_cells = last_row._tr.findall(f'.//{{{W}}}tc')
        if i < len(orig_cells):
            orig_p = orig_cells[i].find(f'{{{W}}}p')
            if orig_p is not None:
                orig_pPr = orig_p.find(f'{{{W}}}pPr')
                if orig_pPr is not None:
                    new_p.append(copy.deepcopy(orig_pPr))
        text = cell_texts[i] if i < len(cell_texts) else ''
        new_p.append(make_run(text))
        tc.append(new_p)
    # Append new row to table
    table._tbl.append(new_tr)


# ============================================================
# MAIN SCRIPT
# ============================================================

print("Backing up original document...")
shutil.copy2(SRC, BACKUP)
print(f"  Backup: {BACKUP}")

print("Loading document...")
doc = Document(SRC)

# ============================================================
# 1. CHAPTER 1.2 – Add two-stage bullet after runtime bullet
# ============================================================
print("\n[1] Updating Chapter 1.2 - Mục tiêu...")

# Find the runtime system bullet (para [112])
runtime_para = find_para_by_text(doc, '• Xây dựng hệ thống runtime hoàn chỉnh')
if runtime_para:
    new_bullet = make_normal_para(
        'Mở rộng hệ thống phân loại từ nhị phân sang kiến trúc hai tầng (two-stage): '
        'tầng 1 (Stage 1) giữ nguyên binary gate Benign/Attack làm production gate chính; '
        'tầng 2 (Stage 2) phân loại họ tấn công (DDoS, DoS, Mirai, Spoofing, Web-Based) '
        'chỉ cho các record đã được tầng 1 đánh dấu Attack; '
        'với cơ chế abstain/unknown cho các họ ngoài phân phối huấn luyện (OOD), '
        'và tích hợp hiển thị thông tin family trên operator console.',
        bullet=True
    )
    insert_after(runtime_para._p, new_bullet)
    print("  ✓ Added two-stage bullet in mục tiêu")
else:
    print("  ✗ Could not find runtime bullet paragraph")

# ============================================================
# 2. CHAPTER 1.3 – Update phạm vi (binary-only limitation)
# ============================================================
print("\n[2] Updating Chapter 1.3 - Phạm vi...")

pham_vi_para = find_para_by_text(doc, 'Về bài toán học máy, đồ án tập trung vào phân loại nhị phân')
if pham_vi_para:
    set_para_text(pham_vi_para,
        'Về bài toán học máy, đồ án xây dựng hệ thống phân loại hai tầng (two-stage): '
        'tầng 1 là mô hình nhị phân Benign/Attack đóng vai trò production gate chính, '
        'và tầng 2 là mô hình phân loại họ tấn công chạy tiếp theo trên các record bị '
        'đánh dấu Attack bởi tầng 1. Tầng 2 hỗ trợ năm họ tấn công trong tập đóng '
        '(DDoS, DoS, Mirai, Spoofing, Web-Based) và có cơ chế abstain để trả về nhãn '
        '"unknown" khi gặp các họ tấn công ngoài phân phối huấn luyện (OOD) mà mô hình '
        'chưa học. Thiết kế này bảo toàn binary alert contract hiện có trong khi bổ sung '
        'thêm thông tin enrichment về loại tấn công một cách additive.')
    print("  ✓ Updated phạm vi paragraph")
else:
    print("  ✗ Could not find phạm vi paragraph")

# ============================================================
# 3. CHAPTER 1.5 – Update Chapter 5 description
# ============================================================
print("\n[3] Updating Chapter 1.5 - Bố cục (Chapter 5 desc)...")

chuong5_desc = find_para_by_text(doc, '• Chương 5 — Triển khai hệ thống: Mô tả chi tiết')
if chuong5_desc:
    set_para_text(chuong5_desc,
        '• Chương 5 — Triển khai hệ thống: Mô tả chi tiết từng thành phần runtime '
        'gồm model bundle, inference engine, realtime pipeline, record adapter, '
        'live sensor, operator console, và hệ thống phân loại hai tầng '
        '(composite bundle, Stage 2 family classifier, cơ chế abstain/unknown, '
        'và tích hợp hiển thị family trên operator console).')
    print("  ✓ Updated Chapter 5 description in bố cục")
else:
    print("  ✗ Could not find Chapter 5 bố cục paragraph")

# ============================================================
# 4. INSERT SECTION 5.7 after end of 5.6
# ============================================================
print("\n[4] Adding Section 5.7 - Hệ thống phân loại hai tầng...")

# Find the last paragraph of 5.6 (about vòng đời schema)
vongle_para = find_para_by_text(doc,
    'Vòng đời schema của console cũng được thiết kế explicit')
if vongle_para:
    # Build section 5.7 paragraphs (insert in REVERSE order since each goes right after vongle_para)
    paras_57 = []

    # Heading
    paras_57.append(make_heading2('5.7 Hệ thống phân loại hai tầng (Two-Stage Classification)'))

    # Intro
    paras_57.append(make_normal_para(
        'Hệ thống phân loại hai tầng là thành phần mở rộng được bổ sung sau giai đoạn '
        'phát triển ban đầu, nâng cấp từ phân loại nhị phân sang kiến trúc hai giai đoạn. '
        'Thiết kế này tuân theo nguyên tắc additive enrichment: tầng 1 (Stage 1) giữ '
        'nguyên binary gate Benign/Attack làm production gate chính, không thay đổi bất '
        'kỳ contract nào hiện có; tầng 2 (Stage 2) chỉ chạy trên các record bị đánh dấu '
        'Attack và bổ sung thông tin họ tấn công như một enrichment layer phía trên.'))

    # 5.7.1 heading
    paras_57.append(make_heading3('5.7.1 Kiến trúc và composite bundle'))

    paras_57.append(make_normal_para(
        'Hai tầng được đóng gói trong một composite bundle duy nhất, thay thế legacy binary '
        'bundle. Composite bundle mở rộng cấu trúc của model_bundle.json (manifest_version = 3) '
        'với hai model component: stage1 (binary CatBoost, giữ nguyên từ catboost_full_data_v1) '
        'và stage2 (multiclass CatBoost được huấn luyện riêng trên các attack rows). '
        'Activation record vẫn là file JSON host-local duy nhất; runtime resolve cả hai '
        'tầng từ một activation path. Nếu Stage 2 validation fail khi khởi động, '
        'runtime fail-closed thay vì silently fallback về binary-only mode.'))

    paras_57.append(make_normal_para(
        'Stage 2 được huấn luyện trên các attack-only rows từ tập train, với không gian '
        'nhãn đóng (closed-set) gồm năm họ: DDoS, DoS, Mirai, Spoofing, và Web-Based. '
        'Hai họ BruteForce và Recon bị giữ lại như OOD probe để đánh giá khả năng '
        'generalization. Mô hình Stage 2 là CatBoost multiclass với class_weight_exponent=0.5 '
        'và iterations=500, được huấn luyện trên 18.457.115 attack rows (full-data).'))

    # Table caption for Stage 2 metrics
    paras_57.append(make_normal_para(
        'Bảng 5.4 tổng hợp kết quả của Stage 2 family classifier trên tập test.'))

    # 5.7.2 heading
    paras_57.append(make_heading3('5.7.2 Output contract và cơ chế abstain'))

    paras_57.append(make_normal_para(
        'Khi Stage 1 trả về nhãn Attack và Stage 2 có đủ độ tin cậy, runtime bổ sung '
        'bốn trường mới vào mỗi output record: attack_family (tên họ tấn công hoặc null), '
        'attack_family_confidence (xác suất top-1 của Stage 2, từ 0 đến 1), '
        'attack_family_margin (hiệu giữa top-1 và top-2, phản ánh sự rõ ràng của quyết định), '
        'và family_status (chuỗi enum ba giá trị: "benign", "unknown", "known"). '
        'Tất cả bốn trường binary gốc (attack_score, predicted_label, is_alert, threshold) '
        'giữ nguyên và vẫn là nguồn chân lý chính cho alert workflow.'))

    paras_57.append(make_normal_para(
        'Cơ chế abstain được triển khai bằng hai ngưỡng được calibrate trên tập validation: '
        'top1_confidence phải vượt ngưỡng thứ nhất, và runner_up_margin phải vượt ngưỡng thứ hai. '
        'Nếu không đủ cả hai điều kiện, family_status được set thành "unknown" thay vì ép '
        'record vào một họ đã biết. Cơ chế này đặc biệt quan trọng với các họ OOD như '
        'BruteForce và Recon: thay vì trả về nhãn sai, Stage 2 thừa nhận sự không chắc chắn. '
        'Hành vi "unknown" không có nghĩa là benign — Stage 1 vẫn đã đánh dấu record đó là '
        'Attack và alert đã được phát.'))

    # 5.7.3 heading
    paras_57.append(make_heading3('5.7.3 Kết quả Stage 2 và so sánh với baseline đa lớp trực tiếp'))

    paras_57.append(make_normal_para(
        'Stage 2 đạt Weighted F1 = 0,9804 và Accuracy = 0,9775 trên tập test, phản ánh '
        'hiệu năng tốt trên các họ tấn công đa số. Tuy nhiên, Macro F1 = 0,5376 cho thấy '
        'sự mất cân bằng đáng kể giữa các họ: DoS đạt F1 = 0,9935 và DDoS đạt F1 = 0,9338, '
        'trong khi Mirai = 0,4488, Spoofing = 0,2955, và Web-Based = 0,0162. '
        'Các họ thiểu số như Web-Based có rất ít dữ liệu trong tập train so với DoS/DDoS, '
        'dẫn đến mô hình thiên lệch mạnh về các họ đa số.'))

    paras_57.append(make_normal_para(
        'So sánh với baseline phân loại đa lớp trực tiếp (Direct Multiclass — dự đoán '
        'Benign + tất cả họ trong một bước): Direct Multiclass đạt Weighted F1 = 0,9688 '
        'và Macro F1 = 0,5008. Thiết kế hai tầng vượt trội ở hai điểm quan trọng: '
        '(1) Macro F1 cao hơn (0,5376 vs 0,5008) nhờ Stage 1 đã lọc sạch Benign trước '
        'khi Stage 2 phân loại; '
        '(2) Direct Multiclass ép 100% lưu lượng OOD vào họ đã biết hoặc Benign, '
        'trong khi hai tầng có khả năng abstain/unknown, phản ánh bản chất open-world '
        'của IDS thực tế.'))

    # 5.7.4 heading
    paras_57.append(make_heading3('5.7.4 Tích hợp trên operator console'))

    paras_57.append(make_normal_para(
        'Thông tin family enrichment được hiển thị trực tiếp trên hai màn hình operator console: '
        'màn hình Alerts (danh sách cảnh báo) hiển thị compact family badge/status cho mỗi alert, '
        'cho phép operator nhận biết họ tấn công ngay khi scan qua danh sách; '
        'màn hình Alert Detail hiển thị đầy đủ family label, family_status, confidence, '
        'margin, và ngữ cảnh abstention để operator hiểu vì sao một cảnh báo được phân loại '
        '"known" hoặc "unknown". Thiết kế UI tuân theo nguyên tắc triage-first: '
        'queue scan nhanh, detail view giải thích sâu. '
        'Quan trọng là "unknown" được trình bày rõ ràng là Stage 1 đã xác nhận tấn công '
        'nhưng Stage 2 không đủ tự tin để gán họ, không phải benign hay lỗi runtime.'))

    # Insert all in reverse order (so they appear in correct order after vongle_para)
    for p_elem in reversed(paras_57):
        insert_after(vongle_para._p, p_elem)

    print(f"  ✓ Inserted {len(paras_57)} paragraphs for Section 5.7")
else:
    print("  ✗ Could not find end of Section 5.6")

# ============================================================
# 5. UPDATE TABLE 16 (Bảng 6.1) - test count
# ============================================================
print("\n[5] Updating Bảng 6.1 - test count...")

table16 = doc.tables[16]
# Row 5: Regression test suite
row5 = table16.rows[5]
cells = row5.cells
# Update cells: "98 tests" → updated count, "98 passed, 0 failed" → updated
if '98 tests' in cells[1].text:
    # Replace paragraph text in cell
    for p in cells[1].paragraphs:
        if p.text:
            set_para_text(p, '127 tests')
    for p in cells[3].paragraphs:
        if p.text:
            set_para_text(p, '127 passed, 0 failed')
    print("  ✓ Updated test count in Bảng 6.1")
else:
    print(f"  ✗ Unexpected cell text: {cells[1].text!r}")

# ============================================================
# 6. UPDATE paragraph about "98 regression test" in Ch 6.1
# ============================================================
print("\n[6] Updating Chapter 6.1 - '98 regression test' text...")

test_para_61 = find_para_by_text(doc, 'Cuối cùng, toàn bộ 98 regression test đều pass')
if test_para_61:
    set_para_text(test_para_61,
        'Cuối cùng, toàn bộ 127 regression test đều pass trong vòng thời gian hợp lý. '
        'Các test này bao phủ các thành phần quan trọng nhất: inference engine, realtime '
        'pipeline, record adapter, feature contract, model bundle manifest, activation '
        'record resolution, và hai tầng phân loại (Stage 1 binary và Stage 2 family '
        'classifier cùng composite bundle contract). Việc 127 test cùng pass là tín hiệu '
        'mạnh rằng hệ thống không có regression ẩn nào tại thời điểm hoàn thành đồ án.')
    print("  ✓ Updated 98→127 in Chapter 6.1 paragraph")
else:
    print("  ✗ Could not find 98-regression-test paragraph in Ch 6.1")

# ============================================================
# 7. UPDATE Chapter 6.2 - Add two-stage bullet after last bullet
# ============================================================
print("\n[7] Adding two-stage bullet in Chapter 6.2 (Điểm mạnh)...")

model_bundle_para = find_para_by_text(doc,
    '• Model bundle versioned: không thể trộn lẫn model/schema/threshold')
if model_bundle_para:
    new_strength = make_normal_para(
        'Phân loại họ tấn công hai tầng: hệ thống không chỉ phát hiện tấn công mà còn '
        'phân loại được họ tấn công (DDoS, DoS, Mirai, Spoofing, Web-Based) với '
        'Weighted F1 = 0,9804, vượt trội so với baseline phân loại đa lớp trực tiếp '
        '(0,9688). Cơ chế abstain/unknown đảm bảo mô hình không ép traffic OOD vào '
        'họ sai, phù hợp với môi trường IDS thực tế. Thông tin family được tích hợp '
        'trực tiếp trên operator console, nâng cao chất lượng triage cho analyst.',
        bullet=True
    )
    insert_after(model_bundle_para._p, new_strength)
    print("  ✓ Added two-stage strength bullet")
else:
    print("  ✗ Could not find model_bundle bullet")

# ============================================================
# 8. UPDATE Chapter 6.2 - "98 regression test" in strengths bullet
# ============================================================
print("\n[8] Updating '98 regression test' bullet in Chapter 6.2...")

coverage_para = find_para_by_text(doc, '• Coverage test rộng: 98 regression test')
if coverage_para:
    set_para_text(coverage_para,
        '• Coverage test rộng: 127 regression test bao phủ các thành phần then chốt '
        'bao gồm cả Stage 1/Stage 2 và composite bundle contract, '
        'cung cấp safety net quan trọng cho các thay đổi trong tương lai.')
    print("  ✓ Updated 98→127 in coverage bullet")
else:
    print("  ✗ Could not find coverage test bullet")

# ============================================================
# 9. UPDATE Chapter 6.3 - Add Stage 2 quality limitation bullet
# ============================================================
print("\n[9] Adding Stage 2 quality limitation in Chapter 6.3...")

# Insert after CICFlowMeter dependency bullet (last bullet in 6.3)
live_sensor_para = find_para_by_text(doc,
    '• Live sensor phụ thuộc vào CICFlowMeter')
if live_sensor_para:
    new_weakness = make_normal_para(
        'Chất lượng tầng 2 không đồng đều giữa các họ tấn công: Macro F1 = 0,5376 '
        'che giấu sự mất cân bằng nghiêm trọng — DoS đạt F1 = 0,9935 và DDoS = 0,9338, '
        'trong khi Mirai = 0,4488, Spoofing = 0,2955, và Web-Based chỉ đạt 0,0162. '
        'Nguyên nhân chính là class imbalance cực đoan giữa các họ trong tập train '
        '(DoS/DDoS chiếm tỷ lệ áp đảo so với Spoofing/Web-Based). '
        'Cải thiện các kỹ thuật xử lý class imbalance và thu thập thêm dữ liệu '
        'cho các họ thiểu số là ưu tiên hàng đầu cho phiên bản tiếp theo.',
        bullet=True
    )
    insert_after(live_sensor_para._p, new_weakness)
    print("  ✓ Added Stage 2 quality limitation bullet")
else:
    print("  ✗ Could not find CICFlowMeter dependency bullet")

# ============================================================
# 10. UPDATE TABLE 17 (Bảng 6.2) - Add two-stage row + update test count row
# ============================================================
print("\n[10] Updating Bảng 6.2 - adding two-stage row...")

table17 = doc.tables[17]
# First update the last row ("Đánh giá hiệu năng hệ thống" -> update test count)
last_data_row = table17.rows[-1]
if 'Đánh giá hiệu năng' in last_data_row.cells[0].text:
    # Update result cell: 98 tests → 127 tests
    result_cell = last_data_row.cells[2]
    for p in result_cell.paragraphs:
        if '98 tests' in p.text:
            set_para_text(p, '127 tests passed, đo throughput và latency end-to-end')
    print("  ✓ Updated test count in Bảng 6.2 last row")

# Add new row for two-stage
add_table_row(table17, [
    'Phân loại họ tấn công hai tầng (Stage 2)',
    'Hoàn thành đầy đủ',
    'Composite bundle, 5 họ, Weighted F1=0,9804, cơ chế abstain/unknown, hiển thị trên console'
])
print("  ✓ Added two-stage row in Bảng 6.2")

# ============================================================
# 11. UPDATE Chapter 7.1 – Add two-stage to conclusion
# ============================================================
print("\n[11] Updating Chapter 7.1 - conclusion...")

test_98_para_71 = find_para_by_text(doc,
    'Với 98 regression test đều pass và toàn bộ end-to-end pipeline')
if test_98_para_71:
    set_para_text(test_98_para_71,
        'Với 127 regression test đều pass và toàn bộ end-to-end pipeline hoạt động '
        'thông suốt, hệ thống đã vượt qua mức "demo" đơn thuần và đạt mức có thể '
        'dùng làm nền tảng cho giai đoạn phát triển tiếp theo thành một sản phẩm IDS '
        'thực sự. Hệ thống phân loại hai tầng — với Stage 1 là binary gate đã được '
        'chứng minh qua thực nghiệm và Stage 2 là family classifier Weighted F1 = 0,9804 '
        '— là bước đi quan trọng từ "phát hiện tấn công" sang "hiểu loại tấn công", '
        'cung cấp thông tin enrichment có giá trị thực tiễn cho SOC analyst. '
        'Các artifact của đồ án, bao gồm mã nguồn, dữ liệu đã xử lý, composite model '
        'bundle, scripts huấn luyện và runbook vận hành, đều được lưu đầy đủ và có '
        'khả năng tái lập.')
    print("  ✓ Updated conclusion paragraph")
else:
    print("  ✗ Could not find '98 regression test' in Chapter 7.1")

# ============================================================
# 12. UPDATE Chapter 7.2 – Change Hướng thứ nhất (multiclass is done)
# ============================================================
print("\n[12] Updating Chapter 7.2 - Hướng thứ nhất...")

huong_1_para = find_para_by_text(doc,
    'Hướng thứ nhất — Phân loại đa lớp theo họ tấn công')
if huong_1_para:
    set_para_text(huong_1_para,
        'Hướng thứ nhất — Hoàn thiện hệ thống phân loại hai tầng đã xây dựng: '
        'Đồ án đã triển khai thành công kiến trúc hai tầng với Stage 2 family classifier '
        'đạt Weighted F1 = 0,9804. Tuy nhiên, Macro F1 = 0,5376 cho thấy hiệu năng '
        'chưa đồng đều giữa các họ tấn công, đặc biệt Spoofing (F1 = 0,2955) và '
        'Web-Based (F1 = 0,0162) cần cải thiện đáng kể. '
        'Hướng phát triển gồm: (a) nghiên cứu các kỹ thuật xử lý class imbalance nâng '
        'cao như focal loss, SMOTE cho dữ liệu tabular, hay cost-sensitive learning; '
        '(b) thu thập thêm dữ liệu cho các họ thiểu số; '
        '(c) đánh giá Stage 2 trên traffic thực tế ngoài phạm vi bộ dữ liệu CIC IoT-DIAD; '
        '(d) mở rộng thêm các họ tấn công mới vào không gian nhãn của Stage 2.')
    print("  ✓ Updated Hướng thứ nhất")
else:
    print("  ✗ Could not find Hướng thứ nhất paragraph")

# ============================================================
# 13. UPDATE Chapter 7.2 – renumber remaining directions
# ============================================================
print("\n[13] Updating Chapter 7.2 - update direction numbering intro...")

huong_intro = find_para_by_text(doc,
    'Dựa trên các giới hạn đã nhận diện ở Chương 6, đồ án đề xuất sáu hướng')
if huong_intro:
    # Hướng thứ nhất is now "hoàn thiện hai tầng" (still 6 directions total)
    print("  ✓ (Direction count remains 6, no change needed)")
else:
    print("  ✗ Could not find huong intro paragraph")

# ============================================================
# SAVE DOCUMENT
# ============================================================
print(f"\nSaving updated document to {OUT}...")
doc.save(OUT)
print("Done!")
print(f"\nSummary of changes:")
print("  1. Chapter 1.2: Added two-stage bullet in mục tiêu")
print("  2. Chapter 1.3: Updated phạm vi (no longer binary-only)")
print("  3. Chapter 1.5: Updated Chapter 5 description")
print("  4. Chapter 5:   Added Section 5.7 (two-stage system)")
print("  5. Bảng 6.1:   Updated test count 98→127")
print("  6. Chapter 6.1: Updated test count paragraph")
print("  7. Chapter 6.2: Added two-stage strength bullet")
print("  8. Chapter 6.2: Updated coverage test bullet 98→127")
print("  9. Chapter 6.3: Added Stage 2 quality limitation bullet")
print("  10. Bảng 6.2:   Added two-stage row + updated test count")
print("  11. Chapter 7.1: Updated conclusion with two-stage")
print("  12. Chapter 7.2: Updated Hướng thứ nhất (multiclass now done)")
